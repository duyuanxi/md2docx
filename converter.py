"""Markdown to Word document converter."""

import re
from pathlib import Path

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def _clean_ai_blank_lines(text: str) -> str:
    """Smart clean for AI exports (Doubao/ChatGPT):
    - Remove lines that are only whitespace
    - Merge 3+ consecutive blank lines into 1
    - Trim trailing whitespace from every line
    """
    # Trim trailing whitespace per line
    lines = [line.rstrip() for line in text.split("\n")]
    # Remove lines that became empty after rstrip (whitespace-only lines)
    # but preserve intentional blank lines (single empty strings)
    cleaned = []
    for line in lines:
        cleaned.append(line)
    # Merge 3+ consecutive empty lines → 1 empty line
    result = []
    empty_count = 0
    for line in cleaned:
        if line == "":
            empty_count += 1
            if empty_count <= 1:  # max 1 blank line
                result.append(line)
        else:
            empty_count = 0
            result.append(line)
    return "\n".join(result)


def _set_paragraph_spacing(para, before_pt: float, after_pt: float,
                           line_spacing_pt: float) -> None:
    pf = para.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    pf.line_spacing = Pt(line_spacing_pt)


def _set_run_font(run, font_name: str, size_pt: float, color_hex: str,
                  bold: bool = False, italic: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor.from_string(color_hex.lstrip("#"))
    run.bold = bold
    run.italic = italic
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = r.makeelement(qn("w:rPr"), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def _heading_spacing(level: int, base_line: float) -> tuple[float, float]:
    """Return (space_before, space_after) for a heading level."""
    before_map = {1: 24, 2: 18, 3: 12, 4: 10, 5: 8, 6: 6}
    return (before_map.get(level, 6), 4)


def _apply_heading(doc: Document, level: int, text: str, t: dict) -> None:
    para = doc.add_paragraph()
    prefix = f"h{level}"
    font = t[f"{prefix}_font"]
    size = t[f"{prefix}_size"]
    bold = t[f"{prefix}_bold"]
    color = t[f"{prefix}_color"]

    before, after = _heading_spacing(level, t["line_spacing_pt"])
    _set_paragraph_spacing(para, before, after, t["line_spacing_pt"])

    run = para.add_run(text)
    _set_run_font(run, font, size, color, bold=bold)


def _apply_body_para(doc: Document, text: str, t: dict) -> None:
    if not text:
        return
    para = doc.add_paragraph()
    _set_paragraph_spacing(para, t["para_spacing_before"],
                           t["para_spacing_after"], t["line_spacing_pt"])
    run = para.add_run(text)
    _set_run_font(run, t["body_font"], t["body_size"], t["body_color"])


def _process_inline(element, para, t: dict) -> None:
    for child in element.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                run = para.add_run(text)
                _set_run_font(run, t["body_font"], t["body_size"],
                              t["body_color"])
        elif isinstance(child, Tag):
            tag_name = child.name
            if tag_name in ("strong", "b"):
                run = para.add_run(child.get_text())
                _set_run_font(run, t["body_font"], t["body_size"],
                              t["body_color"], bold=True)
            elif tag_name in ("em", "i"):
                run = para.add_run(child.get_text())
                _set_run_font(run, t["body_font"], t["body_size"],
                              t["body_color"], italic=True)
            elif tag_name == "code":
                run = para.add_run(child.get_text())
                _set_run_font(run, "Courier New", t["body_size"], "#CC0000")
            elif tag_name == "a":
                run = para.add_run(child.get_text())
                _set_run_font(run, t["body_font"], t["body_size"], "#0000FF")
                run.underline = True
            elif tag_name == "br":
                run = para.add_run("\n")
                _set_run_font(run, t["body_font"], t["body_size"],
                              t["body_color"])
            else:
                _process_inline(child, para, t)


def _process_block(element, doc: Document, t: dict) -> None:
    if not isinstance(element, Tag):
        return

    tag = element.name

    if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
        level = int(tag[1])
        _apply_heading(doc, level, element.get_text(strip=True), t)

    elif tag == "p":
        para = doc.add_paragraph()
        _set_paragraph_spacing(para, t["para_spacing_before"],
                               t["para_spacing_after"], t["line_spacing_pt"])
        _process_inline(element, para, t)

    elif tag in ("ul", "ol"):
        for idx, li in enumerate(element.find_all("li", recursive=False)):
            para = doc.add_paragraph()
            pf = para.paragraph_format
            pf.left_indent = Cm(0.6)
            _set_paragraph_spacing(para, 0, 2, t["line_spacing_pt"])

            # Tab stop at the indent position for text after bullet
            tab_stops = pf.tab_stops
            tab_stops.add_tab_stop(Cm(0.6))

            # Bullet marker + tab → text starts at tab stop
            marker = "-" if tag == "ul" else f"{idx + 1}."
            run_marker = para.add_run(marker + "\t")
            _set_run_font(run_marker, t["body_font"], t["body_size"],
                          t["body_color"])
            _process_inline(li, para, t)

    elif tag == "pre":
        code_text = element.get_text().rstrip("\n")
        for line in code_text.split("\n"):
            para = doc.add_paragraph()
            _set_paragraph_spacing(para, 0, 0, t["line_spacing_pt"])
            run = para.add_run(line)
            _set_run_font(run, "Courier New", max(t["body_size"] - 1, 9),
                          "#333333")
            # Light gray background
            pPr = para.paragraph_format._element.get_or_add_pPr()
            shd = pPr.makeelement(qn("w:shd"), {
                qn("w:fill"): "F2F2F2", qn("w:val"): "clear"})
            pPr.append(shd)

    elif tag == "blockquote":
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.left_indent = Cm(1.5)
        _set_paragraph_spacing(para, t["para_spacing_before"],
                               t["para_spacing_after"], t["line_spacing_pt"])
        _process_inline(element, para, t)

    elif tag == "hr":
        para = doc.add_paragraph()
        _set_paragraph_spacing(para, 12, 12, t["line_spacing_pt"])
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("─" * 50)
        _set_run_font(run, t["body_font"], 8, "#999999")

    elif tag == "table":
        rows = element.find_all("tr")
        if not rows:
            return
        num_cols = max(len(r.find_all(["td", "th"])) for r in rows)
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = "Table Grid"
        for i, row in enumerate(rows):
            cells = row.find_all(["td", "th"])
            for j, cell in enumerate(cells):
                cell_text = cell.get_text(strip=True)
                cell_para = table.cell(i, j).paragraphs[0]
                run = cell_para.add_run(cell_text)
                is_header = cell.name == "th"
                _set_run_font(run, t["body_font"], t["body_size"],
                              t["body_color"], bold=is_header)
        # Space after table
        spacer = doc.add_paragraph()
        _set_paragraph_spacing(spacer, 0, 6, t["line_spacing_pt"])


def _process_elements(elements, doc: Document, t: dict) -> None:
    """Process a list of HTML elements, skipping empty NavigableStrings."""
    for element in elements:
        if isinstance(element, NavigableString):
            text = str(element).strip()
            if text:
                _apply_body_para(doc, text, t)
            continue
        _process_block(element, doc, t)


def convert_text(md_text: str, docx_path: str, template: dict,
                 merge_blank: bool = True) -> str:
    if merge_blank:
        md_text = _clean_ai_blank_lines(md_text)

    html = markdown.markdown(
        md_text, extensions=["tables", "fenced_code"])
    soup = BeautifulSoup(html, "html.parser")

    doc = Document()

    # Paper size
    ps = template.get("page_size", "A4")
    _PAGE_DIMS = {"A4": (21.0, 29.7), "A3": (29.7, 42.0),
                  "8K": (26.0, 36.8), "16K": (18.4, 26.0)}
    pw, ph = _PAGE_DIMS.get(ps, _PAGE_DIMS["A4"])

    for section in doc.sections:
        section.page_width = Cm(pw)
        section.page_height = Cm(ph)
        section.top_margin = Cm(template["margin_top_cm"])
        section.bottom_margin = Cm(template["margin_bottom_cm"])
        section.left_margin = Cm(template["margin_left_cm"])
        section.right_margin = Cm(template["margin_right_cm"])

    body = soup.find("body")
    top_elements = list(body.children) if body else list(soup.children)
    _process_elements(top_elements, doc, template)

    doc.save(docx_path)
    return docx_path


def convert(md_path: str, docx_path: str, template: dict,
            merge_blank: bool = True) -> str:
    md_text = Path(md_path).read_text(encoding="utf-8")
    return convert_text(md_text, docx_path, template, merge_blank)
